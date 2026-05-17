#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "converted" / "main_fixed.docx"
FIGURES = ROOT / "figures"
CHAPTERS = ROOT / "chapter"

CHAPTER_ORDER = [
    CHAPTERS / "abstract.tex",
    CHAPTERS / "ch1_intro.tex",
    CHAPTERS / "ch2_requirements.tex",
    CHAPTERS / "ch3_design.tex",
    CHAPTERS / "ch4_implementation.tex",
    CHAPTERS / "ch5_testing.tex",
    CHAPTERS / "ch6_conclusion.tex",
    CHAPTERS / "acknowledgement.tex",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(clean_inline(text))
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size: float = 12) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)


def add_para(doc: Document, text: str, style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=10.5 if style not in {"Title", "Heading 1", "Heading 2", "Heading 3"} else 12)
    return p


def read_tex(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def strip_comments(text: str) -> str:
    lines = []
    for line in text.splitlines():
        if line.lstrip().startswith("%"):
            continue
        lines.append(re.sub(r"(?<!\\)%.*$", "", line))
    return "\n".join(lines)


def brace_content(s: str, start: int) -> tuple[str, int]:
    assert s[start] == "{"
    depth = 0
    out = []
    i = start
    while i < len(s):
        ch = s[i]
        if ch == "{":
            depth += 1
            if depth > 1:
                out.append(ch)
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return "".join(out), i + 1
            out.append(ch)
        else:
            out.append(ch)
        i += 1
    return "".join(out), i


def parse_bib_numbers() -> dict[str, int]:
    bbl = ROOT / "main.bbl"
    if not bbl.exists():
        return {}
    keys = re.findall(r"\\bibitem(?:\[[^\]]*\])?\{([^}]+)\}", bbl.read_text(encoding="utf-8"))
    return {key: idx + 1 for idx, key in enumerate(keys)}


BIB_NUMBERS = parse_bib_numbers()


def cite_text(keys: str) -> str:
    nums = []
    for key in keys.split(","):
        key = key.strip()
        nums.append(str(BIB_NUMBERS.get(key, key)))
    return "[" + ",".join(nums) + "]"


def clean_inline(text: str) -> str:
    text = text.replace("\r", "")
    text = re.sub(r"\\zihao\{[^}]+\}", "", text)
    text = re.sub(r"\\upcite\{([^}]+)\}", lambda m: cite_text(m.group(1)), text)
    text = re.sub(r"\\cite\{([^}]+)\}", lambda m: cite_text(m.group(1)), text)
    text = re.sub(r"\\ref\{([^}]+)\}", lambda m: REF_NUMBERS.get(m.group(1), ""), text)
    text = re.sub(r"\\label\{[^}]+\}", "", text)
    text = re.sub(r"\\addcontentsline\{[^}]+\}\{[^}]+\}\{[^}]+\}", "", text)
    text = re.sub(r"\\(texttt|textbf|emph|heiti|songti|bfseries|zihao|footnotesize)\*?(?:\[[^\]]*\])?\{([^{}]*)\}", r"\2", text)
    text = re.sub(r"\\url\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\doi\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    replacements = {
        r"\%": "%",
        r"\_": "_",
        r"\&": "&",
        r"\#": "#",
        r"\{": "{",
        r"\}": "}",
        r"\,": " ",
        r"\quad": " ",
        r"\textwidth": "",
        r"\allowbreak": "",
        r"\rightarrow": "->",
        r"\times": "x",
        "~": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s+([，。；：、）])", r"\1", text)
    text = re.sub(r"([（])\s+", r"\1", text)
    return text.strip()


def collect_refs() -> dict[str, str]:
    refs: dict[str, str] = {}
    fig_no = 0
    tab_no = 0
    for path in CHAPTER_ORDER:
        if not path.exists():
            continue
        text = strip_comments(read_tex(path))
        for env, body in iter_envs(text):
            if env == "figure":
                fig_no += 1
                labels = re.findall(r"\\label\{([^}]+)\}", body)
                for label in labels:
                    refs[label] = str(fig_no)
            if env == "table":
                tab_no += 1
                labels = re.findall(r"\\label\{([^}]+)\}", body)
                for label in labels:
                    refs[label] = str(tab_no)
    return refs


def iter_envs(text: str):
    pattern = re.compile(r"\\begin\{(figure|table|lstlisting|itemize|abstractzh|abstracten)\}(?:\[[^\]]*\])?")
    pos = 0
    while True:
        m = pattern.search(text, pos)
        if not m:
            break
        env = m.group(1)
        start = m.end()
        end_pat = re.compile(r"\\end\{" + re.escape(env) + r"\}")
        e = end_pat.search(text, start)
        if not e:
            break
        yield env, text[start:e.start()]
        pos = e.end()


REF_NUMBERS = collect_refs()


def extract_first_command_arg(body: str, command: str) -> str:
    m = re.search(r"\\" + re.escape(command) + r"(?:\[[^\]]*\])?\{", body)
    if not m:
        return ""
    return brace_content(body, m.end() - 1)[0]


def extract_all_command_args(body: str, command: str) -> list[str]:
    args = []
    pos = 0
    regex = re.compile(r"\\" + re.escape(command) + r"(?:\[[^\]]*\])?\{")
    while True:
        m = regex.search(body, pos)
        if not m:
            return args
        arg, pos = brace_content(body, m.end() - 1)
        args.append(arg)


def add_picture(doc: Document, image_ref: str, width_ratio: float = 0.86) -> None:
    image_ref = image_ref.strip()
    path = ROOT / image_ref
    if not path.exists():
        path = FIGURES / Path(image_ref).name
    if not path.exists():
        add_para(doc, f"[图片缺失：{image_ref}]", align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    max_width = Inches(6.1 * width_ratio / 0.86)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=max_width)


def parse_width_ratio(option: str) -> float:
    m = re.search(r"width\s*=\s*([0-9.]+)\\textwidth", option)
    if m:
        return min(max(float(m.group(1)), 0.35), 0.95)
    return 0.86


def add_caption(doc: Document, prefix: str, no: int, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{prefix} {no} {clean_inline(caption)}")
    set_run_font(r, size=10.5)
    r.bold = True


def add_code_block(doc: Document, code: str, caption: str, no: int) -> None:
    add_caption(doc, "代码", no, caption)
    for raw in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.right_indent = Cm(0.25)
        p.paragraph_format.space_after = Pt(0)
        set_para_shading(p, "F2F2F2")
        r = p.add_run(raw.replace("\t", "    "))
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8.5)


def set_para_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def split_latex_row(row: str) -> list[str]:
    cells = []
    cur = []
    depth = 0
    i = 0
    while i < len(row):
        ch = row[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth = max(0, depth - 1)
        if ch == "&" and depth == 0:
            cells.append("".join(cur))
            cur = []
        else:
            cur.append(ch)
        i += 1
    cells.append("".join(cur))
    return [clean_inline(c) for c in cells]


def add_table(doc: Document, body: str, caption: str, no: int) -> None:
    rows = []
    body = re.sub(r"\\(toprule|midrule|bottomrule|centering|small)\b", "", body)
    body = re.sub(r"\\caption\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}", "", body)
    body = re.sub(r"\\label\{[^}]+\}", "", body)
    body = re.sub(r"\\begin\{tabularx?\}\{[^}]+\}(?:\{[^}]+\})?", "", body)
    body = re.sub(r"\\end\{tabularx?\}", "", body)
    for part in re.split(r"\\\\", body):
        part = part.strip()
        if not part or part.startswith("\\"):
            continue
        if "&" not in part:
            continue
        cells = split_latex_row(part)
        if any(cells):
            rows.append(cells)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    add_caption(doc, "表", no, caption)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            set_cell_text(table.cell(i, j), row[j] if j < len(row) else "", bold=i == 0)
            if i == 0:
                set_cell_shading(table.cell(i, j), "EDEDED")


def add_itemize(doc: Document, body: str) -> None:
    items = re.split(r"\\item", body)
    for item in items:
        item = clean_inline(item)
        if item:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(item)
            set_run_font(r, size=10.5)


def process_text_segment(doc: Document, segment: str) -> None:
    segment = strip_comments(segment)
    pos = 0
    command = re.compile(r"\\(chapter\*?|section\*?|subsection\*?|acknowledgement)\s*(?:\{)?")
    while True:
        m = command.search(segment, pos)
        if not m:
            add_plain_text(doc, segment[pos:])
            return
        add_plain_text(doc, segment[pos:m.start()])
        cmd = m.group(1)
        if cmd == "acknowledgement":
            add_heading(doc, "致谢", 1)
            pos = m.end()
            continue
        if segment[m.end() - 1] == "{":
            title, pos = brace_content(segment, m.end() - 1)
            level = 1 if cmd.startswith("chapter") else 2 if cmd.startswith("section") else 3
            add_heading(doc, clean_inline(title), level)
        else:
            pos = m.end()


def add_heading(doc: Document, text: str, level: int) -> None:
    if not text:
        return
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    set_run_font(r, east_asia="黑体", ascii_font="Times New Roman", size=16 if level == 1 else 14 if level == 2 else 12)


def add_plain_text(doc: Document, text: str) -> None:
    text = re.sub(r"\\clearpage|\\newpage|\\phantomsection", "", text)
    text = re.sub(r"\\keywordszh\{([^}]*)\}", r"关键词：\1", text)
    text = re.sub(r"\\keywordsen\{([^}]*)\}", r"Keywords: \1", text)
    text = re.sub(r"\\noindent", "", text)
    text = re.sub(r"\\begin\{[^}]+\}|\\end\{[^}]+\}", "", text)
    parts = re.split(r"\n\s*\n+", text)
    for part in parts:
        line = clean_inline(part.replace("\n", " "))
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.35
        r = p.add_run(line)
        set_run_font(r, size=10.5)


def process_figure(doc: Document, body: str, fig_no: int) -> None:
    captions = extract_all_command_args(body, "caption")
    caption = captions[-1] if captions else ""
    image_args = extract_all_command_args(body, "includegraphics")
    options = re.findall(r"\\includegraphics\[([^\]]*)\]", body)
    if len(image_args) > 1:
        for idx, img in enumerate(image_args):
            ratio = parse_width_ratio(options[idx]) if idx < len(options) else 0.47
            add_picture(doc, img, min(ratio, 0.47))
            if idx < max(0, len(captions) - 1):
                add_para(doc, f"（{chr(97 + idx)}）{clean_inline(captions[idx])}", align=WD_ALIGN_PARAGRAPH.CENTER)
    elif image_args:
        ratio = parse_width_ratio(options[0]) if options else 0.86
        add_picture(doc, image_args[0], ratio)
    if caption:
        add_caption(doc, "图", fig_no, caption)


def process_document_body(doc: Document, text: str, counters: dict[str, int]) -> None:
    text = strip_comments(text)
    pattern = re.compile(r"\\begin\{(figure|table|lstlisting|itemize|abstractzh|abstracten)\}(?:\[[^\]]*\])?")
    pos = 0
    while True:
        m = pattern.search(text, pos)
        if not m:
            process_text_segment(doc, text[pos:])
            break
        process_text_segment(doc, text[pos:m.start()])
        env = m.group(1)
        end = re.search(r"\\end\{" + re.escape(env) + r"\}", text[m.end():])
        if not end:
            break
        body_start = m.end()
        body_end = m.end() + end.start()
        body = text[body_start:body_end]
        if env == "figure":
            counters["figure"] += 1
            process_figure(doc, body, counters["figure"])
        elif env == "table":
            counters["table"] += 1
            add_table(doc, body, extract_first_command_arg(body, "caption"), counters["table"])
        elif env == "lstlisting":
            counters["code"] += 1
            caption_match = re.search(r"caption=\{([^}]*)\}", m.group(0))
            caption = caption_match.group(1) if caption_match else "代码片段"
            add_code_block(doc, body, caption, counters["code"])
        elif env == "itemize":
            add_itemize(doc, body)
        elif env in {"abstractzh", "abstracten"}:
            add_heading(doc, "摘要" if env == "abstractzh" else "Abstract", 1)
            add_plain_text(doc, body)
        pos = body_end + len(f"\\end{{{env}}}")


def add_reference_section(doc: Document) -> None:
    bbl = ROOT / "main.bbl"
    if not bbl.exists():
        return
    add_heading(doc, "参考文献", 1)
    content = bbl.read_text(encoding="utf-8")
    entries = re.split(r"\\bibitem(?:\[[^\]]*\])?\{[^}]+\}", content)[1:]
    for idx, entry in enumerate(entries, start=1):
        entry = re.sub(r"\\newblock", " ", entry)
        entry = re.sub(r"\\allowbreak", "", entry)
        entry = clean_inline(entry)
        if entry:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.75)
            r = p.add_run(f"[{idx}] {entry}")
            set_run_font(r, size=9)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)


def add_cover_pages(doc: Document) -> None:
    for name in ["conver0.png", "conver1.png", "conver2.png"]:
        path = FIGURES / name
        if path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(path), width=Inches(6.2))
            doc.add_section(WD_SECTION.NEW_PAGE)


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_cover_pages(doc)
    counters = {"figure": 0, "table": 0, "code": 0}
    for path in CHAPTER_ORDER:
        if not path.exists():
            continue
        process_document_body(doc, read_tex(path), counters)
    add_reference_section(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)
    print(f"figures={counters['figure']} tables={counters['table']} code_blocks={counters['code']}")


if __name__ == "__main__":
    main()
