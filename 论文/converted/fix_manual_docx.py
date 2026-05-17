#!/usr/bin/env python3
from __future__ import annotations

import shutil
import tempfile
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree
import re


BASE_DIR = Path(__file__).resolve().parent
ROOT = BASE_DIR.parent
SRC = BASE_DIR / "main (1).docx"
OUT = BASE_DIR / "main_100pct_fixed.docx"

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def repack_docx(src_dir: Path, out: Path) -> None:
    with ZipFile(out, "w", ZIP_DEFLATED) as z:
        for path in sorted(src_dir.rglob("*")):
            if path.is_file():
                z.write(path, path.relative_to(src_dir).as_posix())


def ensure_centered_image_paragraph(anchor) -> None:
    p = anchor
    while p is not None and p.tag != f"{{{NS['w']}}}p":
        p = p.getparent()
    if p is None:
        return
    p_pr = p.find("w:pPr", NS)
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p.insert(0, p_pr)
    ind = p_pr.find("w:ind", NS)
    if ind is not None:
        p_pr.remove(ind)
    jc = p_pr.find("w:jc", NS)
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), "center")


def convert_anchor_to_inline(document_xml: Path) -> int:
    tree = etree.parse(str(document_xml))
    root = tree.getroot()
    anchors = root.xpath(".//wp:anchor", namespaces=NS)
    for anchor in anchors:
        ensure_centered_image_paragraph(anchor)
        inline = etree.Element(f"{{{NS['wp']}}}inline", nsmap=anchor.nsmap)
        for attr in ("distT", "distB", "distL", "distR"):
            inline.set(attr, anchor.get(attr, "0"))
        for child_name in ("extent", "effectExtent", "docPr", "cNvGraphicFramePr"):
            child = anchor.find(f"wp:{child_name}", NS)
            if child is not None:
                inline.append(child)
        graphic = anchor.find("{http://schemas.openxmlformats.org/drawingml/2006/main}graphic")
        if graphic is not None:
            inline.append(graphic)
        anchor.getparent().replace(anchor, inline)
    tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True, standalone=True)
    return len(anchors)


def normalize_image_paragraphs(document_xml: Path) -> int:
    tree = etree.parse(str(document_xml))
    root = tree.getroot()
    changed = 0
    for drawing in root.xpath(".//w:drawing", namespaces=NS):
        p = drawing
        while p is not None and p.tag != f"{{{NS['w']}}}p":
            p = p.getparent()
        if p is None:
            continue
        for pos in p.xpath(".//w:rPr/w:position", namespaces=NS):
            pos.getparent().remove(pos)
            changed += 1
        p_pr = p.find("w:pPr", NS)
        if p_pr is not None:
            spacing = p_pr.find("w:spacing", NS)
            if spacing is not None:
                for key in ("line", "lineRule", "before", "after"):
                    qname = f"{{{NS['w']}}}{key}"
                    if qname in spacing.attrib:
                        del spacing.attrib[qname]
                        changed += 1
    tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True, standalone=True)
    return changed


def _shift_style_number(style: str, key: str, delta: float) -> tuple[str, bool]:
    pattern = re.compile(rf"({re.escape(key)}:)(-?\d+(?:\.\d+)?)(pt)?")
    def repl(match):
        value = float(match.group(2)) + delta
        value_text = f"{value:.4f}".rstrip("0").rstrip(".")
        return f"{match.group(1)}{value_text}{match.group(3) or ''}"
    new_style, count = pattern.subn(repl, style, count=1)
    return new_style, count > 0


def _w_el(tag: str, text: str | None = None, **attrs):
    el = OxmlElement(f"w:{tag}")
    for key, val in attrs.items():
        el.set(qn(f"w:{key}"), str(val))
    if text is not None:
        el.text = text
    return el


def _set_text_run(run, text: str, font: str = "Times New Roman", east_asia: str = "宋体",
                  size: int = 21, bold: bool = False) -> None:
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), font)
    r_pr.append(r_fonts)
    r_pr.append(_w_el("sz", val=size))
    r_pr.append(_w_el("szCs", val=size))
    if bold:
        r_pr.append(OxmlElement("w:b"))
        r_pr.append(OxmlElement("w:bCs"))
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    run.append(t)


def make_caption_paragraph(text: str):
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_pr.append(_w_el("jc", val="center"))
    p_pr.append(_w_el("spacing", before="120", after="80", line="300", lineRule="exact"))
    p.append(p_pr)
    r = OxmlElement("w:r")
    _set_text_run(r, text, size=21, bold=True)
    p.append(r)
    return p


def _make_cell(text: str, width: int, shaded: bool = False, align: str = "left", code: bool = False):
    tc = OxmlElement("w:tc")
    tc_pr = OxmlElement("w:tcPr")
    tc_pr.append(_w_el("tcW", w=width, type="dxa"))
    if shaded:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F8FAFC")
        tc_pr.append(shd)
    tc.append(tc_pr)
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_pr.append(_w_el("spacing", before="0", after="0", line="210", lineRule="exact"))
    p_pr.append(_w_el("jc", val=align))
    p.append(p_pr)
    r = OxmlElement("w:r")
    _set_text_run(r, text, font="Consolas" if code else "Times New Roman",
                  east_asia="Consolas" if code else "宋体", size=16 if code else 16)
    p.append(r)
    tc.append(p)
    return tc


def make_code_table(code: str):
    lines = code.strip("\n").splitlines() or [""]
    tbl = OxmlElement("w:tbl")
    tbl_pr = OxmlElement("w:tblPr")
    tbl_pr.append(_w_el("tblW", w="9000", type="dxa"))
    tbl_pr.append(_w_el("jc", val="center"))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_w_el(edge, val="single", sz="4", space="0", color="D0D7DE"))
    tbl_pr.append(borders)
    tbl_cell_mar = OxmlElement("w:tblCellMar")
    for side, value in (("top", 30), ("left", 60), ("bottom", 30), ("right", 60)):
        tbl_cell_mar.append(_w_el(side, w=value, type="dxa"))
    tbl_pr.append(tbl_cell_mar)
    tbl.append(tbl_pr)
    grid = OxmlElement("w:tblGrid")
    grid.append(_w_el("gridCol", w="9000"))
    tbl.append(grid)
    for line in lines:
        tr = OxmlElement("w:tr")
        tr.append(_make_cell(line, 9000, shaded=False, align="left", code=True))
        tbl.append(tr)
    return tbl


def extract_latex_listings() -> list[tuple[str, str]]:
    listings: list[tuple[str, str]] = []
    chapter_files = [
        ROOT / "chapter" / "ch4_implementation.tex",
        ROOT / "chapter" / "ch5_testing.tex",
    ]
    for path in chapter_files:
        text = path.read_text(encoding="utf-8")
        for m in re.finditer(r"\\begin\{lstlisting\}\[[^\]]*caption=\{([^}]*)\}[^\]]*\](.*?)\\end\{lstlisting\}", text, re.S):
            listings.append((m.group(1), m.group(2).strip("\n")))
    return listings


def listing_number(index: int) -> str:
    return f"4.{index}" if index <= 15 else f"5.{index - 15}"


def element_text(el) -> str:
    return "".join(el.xpath(".//w:t/text()", namespaces=NS)).strip()


def is_old_code_element(text: str, tag_name: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return True
    if compact.isdigit():
        return True
    if compact.startswith("代码"):
        return True
    code_tokens = (
        "public", "private", "return", "docker", "curl", "@", "RLock", "SearchRequest",
        "Transactional", "compose", "Map<", "List<", "Set<", "Queue<", "try", "catch",
        "throw", "if(", "if ", "else", "//", "cp.env", "InputStream", "validate", "User "
    )
    if any(token in text for token in code_tokens):
        return True
    if tag_name == "tbl" and len(text) > 0 and not re.search(r"[\u4e00-\u9fff]{6,}", text):
        return True
    return False


def replace_code_blocks(document_xml: Path) -> int:
    listings = extract_latex_listings()
    tree = etree.parse(str(document_xml))
    root = tree.getroot()
    body = root.find("w:body", NS)
    replaced = 0
    pos = 0
    for idx, (caption, code) in enumerate(listings, start=1):
        number = listing_number(idx)
        marker_re = re.compile(rf"代码\s*{re.escape(number)}")
        start = None
        for j in range(pos, len(body)):
            text = element_text(body[j])
            if marker_re.search(text):
                start = j
                break
        if start is None:
            continue
        end = start + 1
        while end < len(body):
            text = element_text(body[end])
            tag_name = etree.QName(body[end]).localname
            if not is_old_code_element(text, tag_name):
                break
            end += 1
        display_caption = f"代码 {number}    {caption}"
        caption_p = make_caption_paragraph(display_caption)
        code_tbl = make_code_table(code)
        for _ in range(end - start):
            body.remove(body[start])
        body.insert(start, caption_p)
        body.insert(start + 1, code_tbl)
        pos = start + 2
        replaced += 1
    tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True, standalone=True)
    return replaced


def remove_orphan_line_number_artifacts(document_xml: Path) -> int:
    tree = etree.parse(str(document_xml))
    root = tree.getroot()
    body = root.find("w:body", NS)
    removed = 0
    # Delete leftover line-number paragraphs generated by the PDF->Word tool.
    for el in list(body):
        if etree.QName(el).localname != "p":
            continue
        text = element_text(el)
        compact = re.sub(r"\s+", "", text)
        has_drawing = bool(el.xpath(".//w:drawing|.//w:pict", namespaces=NS))
        has_sect = bool(el.xpath(".//w:pPr/w:sectPr", namespaces=NS))
        if compact.isdigit() and 1 <= len(compact) <= 70 and not has_drawing and not has_sect:
            body.remove(el)
            removed += 1
    # Delete leftover VML textboxes that only contain line numbers.
    for tb in root.xpath(".//w:txbxContent", namespaces=NS):
        text = "".join(tb.xpath(".//w:t/text()", namespaces=NS))
        compact = re.sub(r"\s+", "", text)
        if not (compact.isdigit() and 1 <= len(compact) <= 70):
            continue
        pict = tb
        while pict is not None and pict.tag != f"{{{NS['w']}}}pict":
            pict = pict.getparent()
        if pict is not None:
            parent = pict.getparent()
            if parent is not None:
                parent.remove(pict)
                removed += 1
    tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True, standalone=True)
    return removed


def remove_code_adjacent_empty_paragraphs(document_xml: Path) -> int:
    tree = etree.parse(str(document_xml))
    root = tree.getroot()
    body = root.find("w:body", NS)
    children = list(body)
    code_related: set[int] = set()
    for i, el in enumerate(children):
        text = element_text(el)
        if re.match(r"^代码\s+\d+\.\d+", text):
            code_related.add(i)
            if i + 1 < len(children) and etree.QName(children[i + 1]).localname == "tbl":
                code_related.add(i + 1)
    removed = 0
    for i, el in reversed(list(enumerate(children))):
        if etree.QName(el).localname != "p":
            continue
        text = element_text(el)
        has_draw = bool(el.xpath(".//w:drawing|.//w:pict", namespaces=NS))
        if text or has_draw:
            continue
        if any(abs(i - j) <= 8 for j in code_related):
            body.remove(el)
            removed += 1
    tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True, standalone=True)
    return removed


def _next_nonempty_text(body, start: int) -> str:
    for j in range(start + 1, len(body)):
        text = element_text(body[j])
        if text:
            return re.sub(r"\s+", " ", text)
    return ""


def remove_body_empty_section_paragraphs(document_xml: Path) -> int:
    tree = etree.parse(str(document_xml))
    root = tree.getroot()
    body = root.find("w:body", NS)
    chapter_re = re.compile(r"^第\s*\d+\s*章")
    chapter_like_re = re.compile(r"^(第\s*\d+\s*章|智能体编排系统总结与展望|致\s*谢|参考文献)")
    in_main_body = False
    removed = 0
    for el in list(body):
        text = re.sub(r"\s+", " ", element_text(el))
        if chapter_re.match(text):
            in_main_body = True
        if etree.QName(el).localname != "p":
            continue
        has_sect = bool(el.xpath("./w:pPr/w:sectPr", namespaces=NS))
        if not has_sect or text:
            continue
        has_drawing = bool(el.xpath(".//w:drawing|.//w:pict", namespaces=NS))
        if has_drawing or not in_main_body:
            continue
        next_text = _next_nonempty_text(body, list(body).index(el))
        if chapter_like_re.match(next_text):
            continue
        body.remove(el)
        removed += 1
    tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True, standalone=True)
    return removed


def shift_code_textboxes(document_xml: Path) -> int:
    tree = etree.parse(str(document_xml))
    root = tree.getroot()
    changed = 0
    for tb in root.xpath(".//w:txbxContent", namespaces=NS):
        text = "".join(tb.xpath(".//w:t/text()", namespaces=NS))
        if not text.strip():
            continue
        shape = tb
        while shape is not None and shape.tag != "{urn:schemas-microsoft-com:vml}shape":
            shape = shape.getparent()
        if shape is None:
            continue
        style = shape.get("style") or ""
        # Code boxes and their line-number boxes are VML absolute-positioned.
        # Their y offset is too small in WPS/Word, causing captions/body text to
        # overlap the first code lines. Move only VML textbox content downward.
        looks_like_code = any(token in text for token in (
            "public ", "@Override", "@Async", "docker ", "RLock", "SearchRequest",
            "Transactional", "compose", "curl", "return ", "if ", "try ", "catch "
        ))
        looks_like_line_numbers = text.strip().isdigit() and len(text.strip()) <= 50
        if not (looks_like_code or looks_like_line_numbers):
            continue
        if "top:" in style:
            style, ok = _shift_style_number(style, "top", 260)
        elif "margin-top:" in style:
            style, ok = _shift_style_number(style, "margin-top", 13)
        else:
            ok = False
        if ok:
            shape.set("style", style)
            changed += 1
    tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True, standalone=True)
    return changed


def normalize_code_pict_paragraphs(document_xml: Path) -> int:
    tree = etree.parse(str(document_xml))
    root = tree.getroot()
    changed = 0
    for pict in root.xpath(".//w:pict", namespaces=NS):
        p = pict
        while p is not None and p.tag != f"{{{NS['w']}}}p":
            p = p.getparent()
        if p is None:
            continue
        for pos in p.xpath(".//w:rPr/w:position", namespaces=NS):
            pos.getparent().remove(pos)
            changed += 1
        p_pr = p.find("w:pPr", NS)
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            p.insert(0, p_pr)
        spacing = p_pr.find("w:spacing", NS)
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        # Keep the large exact line height that reserves the code-box area, but
        # add a small top gap so the preceding caption cannot overlap the box.
        spacing.set(qn("w:before"), "120")
        changed += 1
    tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True, standalone=True)
    return changed


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str = "EDEDED") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_centered(table, widths_cm: list[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is not None:
        tbl_pr.remove(tbl_ind)
    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)


def delete_column(table, col_idx: int) -> None:
    grid = table._tbl.tblGrid
    if grid is not None and col_idx < len(grid.gridCol_lst):
        grid.remove(grid.gridCol_lst[col_idx])
    for row in table.rows:
        tr = row._tr
        cells = tr.tc_lst
        if col_idx < len(cells):
            tr.remove(cells[col_idx])


def table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def build_table_before(doc: Document, old_table, rows: list[list[str]], widths_cm: list[float]):
    new_table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    new_table.style = old_table.style
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(new_table.cell(i, j), value, bold=i == 0)
            if i == 0:
                shade_cell(new_table.cell(i, j))
    set_table_centered(new_table, widths_cm)
    old_table._tbl.addprevious(new_table._tbl)
    old_table._tbl.getparent().remove(old_table._tbl)
    return new_table


def fix_typical_case_table(doc: Document, table) -> bool:
    text = table_text(table)
    if "TC-AG-" not in text or "代表性功能测试用例" in text:
        return False
    if "编号" not in text or "预期结果" not in text:
        return False
    rows = [
        ["编号", "模块", "测试场景", "预期结果", "结果"],
        ["TC-AG-03", "Agent 管理", "发布 Agent 版本", "生成版本快照，已发布版本指针更新", "通过"],
        ["TC-WF-01", "工作流执行", "启动标准工作流", "SSE 推送连接事件与节点状态事件", "通过"],
        ["TC-WF-03", "工作流执行", "条件分支路由", "正确选中分支，未选中路径标记为 SKIPPED", "通过"],
        ["TC-HR-01", "人工检查点", "BEFORE_EXECUTION 暂停", "推送 workflow_paused，并加入待审批队列", "通过"],
        ["TC-HR-04", "人工检查点", "审批通过并修改输出", "修改值写入上下文，后续节点使用新结果", "通过"],
        ["TC-HR-06", "人工检查点", "并发审批冲突", "后提交请求被拒绝，返回 409 Conflict", "通过"],
        ["TC-KN-02", "知识库", "上传文档并异步处理", "文档状态按 PENDING、PROCESSING、COMPLETED 流转", "通过"],
        ["TC-AU-08", "用户认证", "未携带令牌访问接口", "请求被拦截，返回未认证错误", "通过"],
    ]
    build_table_before(doc, table, rows, [1.45, 1.65, 3.25, 4.05, 0.9])
    return True


def fix_deploy_env_table(doc: Document, table) -> bool:
    text = table_text(table)
    if "PRIMARY_DB_PASSWORD" not in text or "CORS_ALLOWED_ORIGINS" not in text:
        return False
    rows = [
        ["配置项", "说明"],
        ["PRIMARY_DB_PASSWORD", "MySQL 数据库访问密码"],
        ["REDIS_PASSWORD", "Redis 访问密码，用于缓存和分布式锁连接"],
        ["MINIO_ROOT_PASSWORD", "MinIO 管理员密码，用于对象存储访问"],
        ["JWT_SECRET", "JWT 签名密钥，用于用户登录令牌签发与校验"],
        ["MAIL_HOST / MAIL_USERNAME / MAIL_PASSWORD", "邮件验证码发送所需 SMTP 配置"],
        ["CORS_ALLOWED_ORIGINS", "允许访问后端接口的前端域名"],
    ]
    build_table_before(doc, table, rows, [4.2, 8.1])
    return True


def fix_tables(docx_path: Path) -> tuple[int, int, int]:
    doc = Document(str(docx_path))
    centered = 0
    typical_fixed = 0
    env_fixed = 0
    for table in list(doc.tables):
        if fix_typical_case_table(doc, table):
            typical_fixed += 1
            continue
        if fix_deploy_env_table(doc, table):
            env_fixed += 1
            continue
        if len(table.columns) > 1:
            set_table_centered(table)
            centered += 1
    doc.save(str(docx_path))
    return centered, typical_fixed, env_fixed


def tighten_references(docx_path: Path) -> int:
    doc = Document(str(docx_path))
    in_refs = False
    changed = 0
    for p in doc.paragraphs:
        if p.text.strip() == "参考文献":
            in_refs = True
            continue
        if not in_refs or not p.text.strip():
            continue
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(13)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        changed += 1
    doc.save(str(docx_path))
    return changed


def set_paragraph_font(p, size_pt: float, bold: bool, east_asia: str = "黑体") -> None:
    for run in p.runs:
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_single_run_text(p, text: str) -> None:
    p.clear()
    p.add_run(text)


def normalize_headings(docx_path: Path) -> int:
    doc = Document(str(docx_path))
    changed = 0
    chapter_re = re.compile(r"^第\s*(\d+)\s*章\s*(?!为)([^，。；:：]{1,40})$")
    section_re = re.compile(r"^\d+\.\d+\s+")
    subsection_re = re.compile(r"^\d+\.\d+\.\d+\s+")
    unnumbered = {"智能体编排系统总结与展望", "致谢", "参考文献"}
    for p in doc.paragraphs:
        text = re.sub(r"\s+", " ", p.text.strip())
        compact = re.sub(r"\s+", "", p.text.strip())
        if not text:
            continue
        pf = p.paragraph_format
        chapter_match = chapter_re.match(text)
        if chapter_match or compact in unnumbered:
            if chapter_match:
                set_single_run_text(p, f"第 {chapter_match.group(1)} 章 {chapter_match.group(2).strip()}")
            elif compact == "致谢":
                set_single_run_text(p, "致谢")
            elif compact in unnumbered:
                set_single_run_text(p, compact)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.left_indent = None
            pf.first_line_indent = None
            pf.space_before = Pt(18)
            pf.space_after = Pt(18)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(24)
            set_paragraph_font(p, 16, True, "黑体")
            changed += 1
        elif subsection_re.match(text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.left_indent = None
            pf.first_line_indent = None
            pf.space_before = Pt(6)
            pf.space_after = Pt(3)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(18)
            set_paragraph_font(p, 11, True, "黑体")
            changed += 1
        elif section_re.match(text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.left_indent = None
            pf.first_line_indent = None
            pf.space_before = Pt(10)
            pf.space_after = Pt(6)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(20)
            set_paragraph_font(p, 12, True, "黑体")
            changed += 1
    doc.save(str(docx_path))
    return changed


def normalize_code_captions(docx_path: Path) -> int:
    doc = Document(str(docx_path))
    changed = 0
    previous = None
    to_delete = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not re.match(r"^代码\s+\d+\.\d+", text):
            if text:
                previous = p
            continue
        if previous is not None:
            previous.add_run("  ")
            run = previous.add_run(text)
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            previous.paragraph_format.space_after = Pt(2)
            to_delete.append(p)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
        changed += 1
    for p in to_delete:
        p._element.getparent().remove(p._element)
    doc.save(str(docx_path))
    return changed


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        with ZipFile(SRC) as z:
            z.extractall(tmp_dir)
        document_xml = tmp_dir / "word" / "document.xml"
        anchors = convert_anchor_to_inline(document_xml)
        image_paragraph_changes = normalize_image_paragraphs(document_xml)
        code_blocks_replaced = replace_code_blocks(document_xml)
        line_number_artifacts_removed = remove_orphan_line_number_artifacts(document_xml)
        code_empty_paragraphs_removed = remove_code_adjacent_empty_paragraphs(document_xml)
        empty_section_paragraphs_removed = remove_body_empty_section_paragraphs(document_xml)
        repack_docx(tmp_dir, OUT)
    centered, typical_fixed, env_fixed = fix_tables(OUT)
    headings_fixed = normalize_headings(OUT)
    refs_tightened = tighten_references(OUT)
    print(f"written={OUT}")
    print(f"anchors_converted={anchors}")
    print(f"image_paragraph_changes={image_paragraph_changes}")
    print(f"code_blocks_replaced={code_blocks_replaced}")
    print(f"line_number_artifacts_removed={line_number_artifacts_removed}")
    print(f"code_empty_paragraphs_removed={code_empty_paragraphs_removed}")
    print(f"empty_section_paragraphs_removed={empty_section_paragraphs_removed}")
    print(f"tables_centered={centered}")
    print(f"typical_case_tables_fixed={typical_fixed}")
    print(f"deploy_env_tables_fixed={env_fixed}")
    print(f"headings_fixed={headings_fixed}")
    print(f"reference_paragraphs_tightened={refs_tightened}")


if __name__ == "__main__":
    main()
